import streamlit as st
import yfinance as yf
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io
import requests
from docx import Document

st.set_page_config(page_title="SFM Master Terminal", layout="wide")

# Funzione Ricerca Ticker Robusta
def trova_ticker(nome):
    if not nome: return None
    try:
        url = f"https://query2.finance.yahoo.com/v1/finance/search?q={nome}"
        res = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=5)
        return res.json()['quotes'][0]['symbol']
    except: return nome

st.title("🛡️ SFM Master Intelligence Terminal")

# Input con valore predefinito gestito per evitare crash
nome_azienda = st.text_input("Inserisci Azienda per Analisi Profonda", "Tenaris")
ticker = trova_ticker(nome_azienda)

if ticker:
    try:
        stock = yf.Ticker(ticker)
        # Scarico tutto subito per evitare chiamate multiple
        info = stock.info
        cf = stock.cashflow
        hist = stock.history(period="1y")
        
        if not cf.empty and 'Free Cash Flow' in cf.index and not hist.empty:
            st.header(f"Analisi Valutativa: {info.get('longName', ticker)}")
            
            # --- MODELLO DCF 10 ANNI (ESATTO DAL TUO EXCEL) ---
            fcf_base = cf.loc['Free Cash Flow'].iloc[0]
            cash = info.get('totalCash', 0)
            debt = info.get('totalDebt', 0)
            shares = info.get('sharesOutstanding', 1)
            current_p = info.get('currentPrice', 0)

            st.subheader("💎 Proiezione Flussi di Cassa 10 Anni")
            c1, c2, c3 = st.columns(3)
            g_rate = c1.slider("Crescita Annuale (1-10y) %", 0, 30, 10) / 100
            wacc = c2.slider("WACC (Sconto) %", 5, 15, 10) / 100
            p_rate = c3.slider("Crescita Perpetua %", 1.0, 4.0, 2.5) / 100

            # Tabella di marcia 10 anni
            anni = list(range(1, 11))
            fcf_list = []
            pv_list = []
            fcf_step = fcf_base
            for a in anni:
                fcf_step *= (1 + g_rate)
                fcf_list.append(fcf_step)
                pv_list.append(fcf_step / ((1 + wacc) ** a))
            
            # Valore Terminale e Intrinseco
            tv = (fcf_list[-1] * (1 + p_rate)) / (wacc - p_rate)
            pv_tv = tv / ((1 + wacc) ** 10)
            fair_value = (sum(pv_list) + pv_tv + cash - debt) / shares
            upside = ((fair_value / current_p) - 1) * 100

            # --- ANALISI STRATEGICA LONG/SHORT ---
            st.divider()
            # Calcolo RSI per il Timing
            delta = hist['Close'].diff()
            gain = (delta.where(delta > 0, 0)).rolling(14).mean().iloc[-1]
            loss = (-delta.where(delta < 0, 0)).rolling(14).mean().iloc[-1]
            rsi = 100 - (100 / (1 + (gain/loss)))

            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.metric("Fair Value", f"${fair_value:.2f}", f"{upside:.1f}% Upside")
                if upside > 20 and rsi < 45:
                    st.success("🎯 SEGNALE: STRONG BUY (Value + Timing)")
                elif upside < -10 and rsi > 65:
                    st.error("🎯 SEGNALE: STRONG SELL (Overvalued)")
                else:
                    st.warning("🎯 SEGNALE: NEUTRALE / ATTENDERE")
            
            with col_s2:
                st.metric("RSI Tecnico", f"{rsi:.1f}")

            # --- EXPORT PROFESSIONALE (DATI VERI) ---
            st.divider()
            cw, ce = st.columns(2)
            with cw:
                if st.button("📝 Genera Report Word"):
                    doc = Document()
                    doc.add_heading(f"Analisi SFM Master: {info.get('longName')}", 0)
                    doc.add_heading("Valutazione Intrinseca", level=1)
                    doc.add_paragraph(f"Sulla base di una crescita del {g_rate*100}% e un WACC del {wacc*100}%, il valore stimato è ${fair_value:.2f}.")
                    doc.add_heading("Analisi Business", level=1)
                    doc.add_paragraph(info.get('longBusinessSummary', ''))
                    buf = io.BytesIO()
                    doc.save(buf)
                    st.download_button("Scarica Word", buf.getvalue(), f"Report_{ticker}.docx")
            
            with ce:
                output = io.BytesIO()
                df_dcf = pd.DataFrame({'Anno': anni, 'FCF': fcf_list, 'PV': pv_list})
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df_dcf.to_excel(writer, sheet_name='DCF_Model_10Y')
                    pd.DataFrame([info]).to_excel(writer, sheet_name='Info_Generali')
                st.download_button("Scarica Excel", output.getvalue(), f"Modello_{ticker}.xlsx")

            # --- GRAFICO A CANDELE (IN FONDO E FUNZIONANTE) ---
            st.divider()
            st.subheader("📈 Analisi Tecnica e Trend")
            fig = make_subplots(rows=2, cols=1, shared_xaxes=True, row_heights=[0.7, 0.3])
            fig.add_trace(go.Candlestick(x=hist.index, open=hist['Open'], high=hist['High'], low=hist['Low'], close=hist['Close'], name='Price'), row=1, col=1)
            
            # Calcolo RSI per tutto lo storico grafico
            hist_delta = hist['Close'].diff()
            h_gain = (hist_delta.where(hist_delta > 0, 0)).rolling(14).mean()
            h_loss = (-hist_delta.where(hist_delta < 0, 0)).rolling(14).mean()
            hist['RSI_Plot'] = 100 - (100 / (1 + (h_gain/h_loss)))
            
            fig.add_trace(go.Scatter(x=hist.index, y=hist['RSI_Plot'], name='RSI', line=dict(color='orange')), row=2, col=1)
            fig.update_layout(template='plotly_dark', height=700, xaxis_rangeslider_visible=False)
            st.plotly_chart(fig, use_container_width=True)

        else:
            st.warning("Dati di bilancio o storici insufficienti per l'analisi profonda.")

    except Exception as e:
        st.error(f"Errore tecnico: {e}")